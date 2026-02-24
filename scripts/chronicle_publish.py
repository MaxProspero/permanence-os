#!/usr/bin/env python3
"""
Publish latest chronicle report for agent access and operator distribution.

Outputs:
- Stable local files for agents/code (`memory/chronicle/shared`)
- Optional mirror into Google Drive desktop-sync folder
- Optional `.docx` summary export (for Google Docs import)
- Optional SMTP email delivery
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import smtplib
from datetime import datetime, timezone
from email.message import EmailMessage
from pathlib import Path
from typing import Any

from scripts.chronicle_common import BASE_DIR, CHRONICLE_OUTPUT_DIR, ensure_chronicle_dirs, utc_iso

DEFAULT_SHARED_DIR = Path(BASE_DIR) / "memory" / "chronicle" / "shared"


def _latest_report_paths(explicit_json: str | None, explicit_md: str | None) -> tuple[Path, Path]:
    if explicit_json:
        json_path = Path(explicit_json).expanduser()
        md_path = Path(explicit_md).expanduser() if explicit_md else json_path.with_suffix(".md")
        if not json_path.exists():
            raise FileNotFoundError(f"Chronicle JSON not found: {json_path}")
        return json_path, md_path

    if explicit_md:
        md_path = Path(explicit_md).expanduser()
        json_path = md_path.with_suffix(".json")
        if not md_path.exists():
            raise FileNotFoundError(f"Chronicle markdown not found: {md_path}")
        return json_path, md_path

    candidates = sorted(
        CHRONICLE_OUTPUT_DIR.glob("chronicle_report_*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(
            f"No chronicle report found in {CHRONICLE_OUTPUT_DIR}. Run `python cli.py chronicle-report` first."
        )
    json_path = candidates[0]
    md_path = json_path.with_suffix(".md")
    return json_path, md_path


def _load_report(path: Path) -> dict[str, Any]:
    with open(path, "r") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        raise ValueError(f"Unexpected JSON shape in {path}")
    return data


def _build_summary_markdown(report: dict[str, Any], source_json: Path, source_md: Path) -> str:
    totals = report.get("signal_totals") or {}
    direction_events = report.get("direction_events") or []
    issue_events = report.get("issue_events") or []

    lines = [
        "# Chronicle Publish Summary",
        "",
        f"- Published at: {utc_iso()}",
        f"- Report generated: {report.get('generated_at', 'unknown')}",
        f"- Window days: {report.get('days', 'unknown')}",
        f"- Chronicle events: {report.get('events_count', 0)}",
        f"- Git commits: {report.get('commit_count', 0)}",
        "",
        "## Signal Totals",
        f"- Direction hits: {totals.get('direction_hits', 0)}",
        f"- Frustration hits: {totals.get('frustration_hits', 0)}",
        f"- Issue hits: {totals.get('issue_hits', 0)}",
        f"- Log error hits: {totals.get('log_error_hits', 0)}",
        f"- Log warning hits: {totals.get('log_warning_hits', 0)}",
        "",
        "## Direction Events",
    ]

    if direction_events:
        for item in direction_events[-5:]:
            lines.append(f"- {item.get('timestamp', 'unknown')} | {item.get('summary', '')}")
    else:
        lines.append("- none")

    lines.extend(["", "## Friction Events"])
    if issue_events:
        for item in issue_events[-5:]:
            lines.append(f"- {item.get('timestamp', 'unknown')} | {item.get('summary', '')}")
    else:
        lines.append("- none")

    lines.extend(
        [
            "",
            "## Sources",
            f"- Report JSON: {source_json}",
            f"- Report Markdown: {source_md}",
        ]
    )
    return "\n".join(lines) + "\n"


def _write_docx(summary_md: str, path: Path) -> tuple[bool, str]:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False, "python-docx not installed (skip docx export)"

    doc = Document()
    doc.add_heading("Chronicle Publish Summary", level=1)
    for line in summary_md.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(stripped)

    doc.save(path)
    return True, "ok"


def _copy_all(paths: list[Path], target_dir: Path) -> list[Path]:
    target_dir.mkdir(parents=True, exist_ok=True)
    copied: list[Path] = []
    for src in paths:
        if not src.exists():
            continue
        dest = target_dir / src.name
        shutil.copy2(src, dest)
        copied.append(dest)
    return copied


def _send_email(
    recipients: list[str],
    subject: str,
    body: str,
    attachments: list[Path],
    smtp_host: str | None,
    smtp_port: int,
    smtp_user: str | None,
    smtp_password: str | None,
    smtp_from: str | None,
    use_starttls: bool,
) -> str:
    if not recipients:
        return "no recipients provided"
    if not smtp_host:
        return "missing SMTP host"
    if not smtp_from:
        return "missing sender address (PERMANENCE_EMAIL_FROM)"

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = smtp_from
    msg["To"] = ", ".join(recipients)
    msg.set_content(body)

    for path in attachments:
        if not path.exists():
            continue
        data = path.read_bytes()
        if path.suffix.lower() == ".json":
            maintype, subtype = "application", "json"
        elif path.suffix.lower() == ".md":
            maintype, subtype = "text", "markdown"
        elif path.suffix.lower() == ".docx":
            maintype, subtype = (
                "application",
                "vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            maintype, subtype = "application", "octet-stream"
        msg.add_attachment(data, maintype=maintype, subtype=subtype, filename=path.name)

    with smtplib.SMTP(smtp_host, smtp_port, timeout=30) as server:
        if use_starttls:
            server.starttls()
        if smtp_user:
            server.login(smtp_user, smtp_password or "")
        server.send_message(msg)
    return "sent"


def main() -> int:
    parser = argparse.ArgumentParser(description="Publish latest chronicle report for agents + operator channels.")
    parser.add_argument("--report-json", help="Chronicle report JSON path (default: latest in outputs/chronicle)")
    parser.add_argument("--report-md", help="Chronicle report markdown path (optional)")
    parser.add_argument("--output-dir", default=str(DEFAULT_SHARED_DIR), help="Shared output directory")
    parser.add_argument(
        "--drive-dir",
        default=os.getenv("PERMANENCE_CHRONICLE_DRIVE_DIR"),
        help="Optional Google Drive desktop-sync folder destination",
    )
    parser.add_argument("--docx", action="store_true", help="Also export summary as DOCX")
    parser.add_argument("--email-to", action="append", default=[], help="Email recipient (repeatable)")
    parser.add_argument("--email-subject", default="Permanence OS Chronicle Update", help="Email subject")
    parser.add_argument("--smtp-host", default=os.getenv("PERMANENCE_SMTP_HOST"), help="SMTP host")
    parser.add_argument(
        "--smtp-port",
        type=int,
        default=int(os.getenv("PERMANENCE_SMTP_PORT", "587")),
        help="SMTP port",
    )
    parser.add_argument("--smtp-user", default=os.getenv("PERMANENCE_SMTP_USERNAME"), help="SMTP username")
    parser.add_argument("--smtp-password", default=os.getenv("PERMANENCE_SMTP_PASSWORD"), help="SMTP password")
    parser.add_argument("--smtp-from", default=os.getenv("PERMANENCE_EMAIL_FROM"), help="Sender email")
    parser.add_argument("--no-starttls", action="store_true", help="Disable STARTTLS")
    args = parser.parse_args()

    ensure_chronicle_dirs()
    output_dir = Path(args.output_dir).expanduser()
    output_dir.mkdir(parents=True, exist_ok=True)

    source_json, source_md = _latest_report_paths(args.report_json, args.report_md)
    report = _load_report(source_json)
    summary_md = _build_summary_markdown(report, source_json, source_md)

    stable_json = output_dir / "chronicle_latest.json"
    stable_md = output_dir / "chronicle_latest.md"
    summary_path = output_dir / "chronicle_latest_summary.md"
    manifest_path = output_dir / "chronicle_latest_manifest.json"

    shutil.copy2(source_json, stable_json)
    if source_md.exists():
        shutil.copy2(source_md, stable_md)
    else:
        stable_md.write_text(summary_md, encoding="utf-8")

    archived_json = output_dir / source_json.name
    archived_md = output_dir / source_md.name if source_md.name else output_dir / "chronicle_report_latest.md"
    shutil.copy2(source_json, archived_json)
    if source_md.exists():
        shutil.copy2(source_md, archived_md)

    summary_path.write_text(summary_md, encoding="utf-8")

    exported_paths: list[Path] = [stable_json, stable_md, summary_path, archived_json]
    if archived_md.exists():
        exported_paths.append(archived_md)

    docx_status = "skipped"
    docx_path = output_dir / "chronicle_latest_summary.docx"
    if args.docx:
        ok, status = _write_docx(summary_md, docx_path)
        docx_status = status
        if ok and docx_path.exists():
            exported_paths.append(docx_path)

    manifest = {
        "published_at": utc_iso(),
        "source_json": str(source_json),
        "source_md": str(source_md),
        "stable_json": str(stable_json),
        "stable_md": str(stable_md),
        "summary_md": str(summary_path),
        "docx_path": str(docx_path) if docx_path.exists() else None,
        "signal_totals": report.get("signal_totals", {}),
        "events_count": report.get("events_count", 0),
        "commit_count": report.get("commit_count", 0),
    }
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    exported_paths.append(manifest_path)

    drive_copied: list[Path] = []
    if args.drive_dir:
        drive_dir = Path(args.drive_dir).expanduser()
        drive_copied = _copy_all(exported_paths, drive_dir)

    email_status = "skipped"
    if args.email_to:
        body = (
            summary_md
            + "\nShared files:\n"
            + "\n".join(f"- {path}" for path in exported_paths)
        )
        try:
            email_status = _send_email(
                recipients=args.email_to,
                subject=args.email_subject,
                body=body,
                attachments=[summary_path, stable_md, stable_json] + ([docx_path] if docx_path.exists() else []),
                smtp_host=args.smtp_host,
                smtp_port=args.smtp_port,
                smtp_user=args.smtp_user,
                smtp_password=args.smtp_password,
                smtp_from=args.smtp_from,
                use_starttls=not args.no_starttls,
            )
        except Exception as exc:
            email_status = f"failed: {exc}"

    print(f"Chronicle source JSON: {source_json}")
    print(f"Chronicle source MD:   {source_md}")
    print(f"Shared output dir:     {output_dir}")
    print(f"Summary markdown:      {summary_path}")
    print(f"DOCX export:           {docx_status}")
    if args.drive_dir:
        print(f"Drive mirror count:    {len(drive_copied)}")
    print(f"Email status:          {email_status}")
    print(f"Manifest:              {manifest_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
